import os
import sys
import subprocess
from pathlib import Path
from datetime import datetime
import tkinter as tk
from tkinter import messagebox, ttk
from transformers import pipeline
import textwrap
import time

# Auto-install function
def install_and_import(package):
    try:
        __import__(package)
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])

# Install required libraries
install_and_import("requests")
install_and_import("bs4")
install_and_import("pandas")
install_and_import("openpyxl")
install_and_import("python-docx")
install_and_import("transformers")

import requests
from bs4 import BeautifulSoup
import pandas as pd
from docx import Document

# Initialize summarization model (using Google FLAN-T5)
summarizer = pipeline("summarization", model="google/flan-t5-base")

# Enhanced summarization function with chunking
def summarize_text(text, max_length=150, min_length=50):
    max_token_limit = 1024  # Adjust this limit as per model's capability
    
    # Split the text into chunks if it's too long
    wrapped_text = textwrap.wrap(text, width=max_token_limit)
    summaries = []
    
    for chunk in wrapped_text:
        try:
            summary = summarizer(chunk, max_length=max_length, min_length=min_length, do_sample=False)
            summaries.append(summary[0]['summary_text'])
        except Exception as e:
            print(f"Error summarizing text chunk: {e}")
    
    # Combine all summaries into a final summary
    final_summary = ' '.join(summaries)
    return final_summary

# Function to create a directory for saving files
def create_directory(custom_path):
    try:
        main_folder = Path(custom_path) / "web scrapped files"
        main_folder.mkdir(parents=True, exist_ok=True)

        folder_name = datetime.now().strftime("%Y%m%d_%H%M%S")
        folder_path = main_folder / folder_name
        folder_path.mkdir()
        return folder_path
    except FileNotFoundError:
        messagebox.showerror("Error", "Could not find or create the specified path. Please check the folder path and try again.")
        return None

# Function to fetch the content from a URL
def fetch_content(url):
    response = requests.get(url)
    response.raise_for_status()
    return response.content

# Function to parse and save data from the fetched content
def parse_and_save_data(content, folder_path, output_format):
    soup = BeautifulSoup(content, 'html.parser')
    
    # Extract tables if they exist
    tables = soup.find_all('table')
    if tables and output_format in ["Excel", "CSV"]:
        data_frames = pd.read_html(str(tables))
        for i, df in enumerate(data_frames):
            if output_format == "Excel":
                file_path = folder_path / f"table_data_{i+1}.xlsx"
                df.to_excel(file_path, index=False)
            elif output_format == "CSV":
                file_path = folder_path / f"table_data_{i+1}.csv"
                df.to_csv(file_path, index=False)
        messagebox.showinfo("Success", f"Data saved in {output_format} files.")
    else:
        # Extract paragraphs for summarization
        paragraphs = soup.find_all('p')
        if paragraphs:
            doc = Document()
            for para in paragraphs:
                doc.add_paragraph(para.get_text())
            
            # Summarize the paragraphs
            all_text = " ".join([para.get_text() for para in paragraphs])
            summary = summarize_text(all_text)
            
            # Add summary to the Word document
            doc.add_paragraph("Summary:")
            doc.add_paragraph(summary)
            
            # Save the document
            file_path = folder_path / "text_data_with_summary.docx"
            doc.save(file_path)
            messagebox.showinfo("Success", "Data and summary saved in a Word document.")
        else:
            messagebox.showinfo("Info", "No suitable data found on the page.")

# Function to run the web scraping process
def run_scraping(urls, custom_path):
    custom_path = custom_path.strip().strip('"').strip("'")
    folder_path = create_directory(custom_path)
    if folder_path:
        total_urls = len(urls.splitlines())
        for i, url in enumerate(urls.splitlines()):
            try:
                content = fetch_content(url)
                output_format = format_var.get()  # Get selected format
                parse_and_save_data(content, folder_path, output_format)
                
                # Update the progress bar
                progress_bar['value'] = (i + 1) / total_urls * 100
                root.update_idletasks()
                time.sleep(0.2)  # Slight delay for better UI feedback
            except requests.RequestException as e:
                messagebox.showerror("Error", f"Failed to retrieve data from the URL: {e}")

# Function to start scraping when the "Start Scraping" button is pressed
def start_scraping():
    urls = url_entry.get("1.0", tk.END).strip()  # Get text from Text widget
    custom_path = path_entry.get()  # Get the path input
    if not urls or not custom_path:
        messagebox.showerror("Error", "Please provide both URLs and a storage path.")
        return
    
    # Update status label and display a start message
    status_label.config(text="Scraping and summarization have started...", fg="blue")
    root.update_idletasks()  # Ensure the UI updates before processing
    
    # Run the scraping process
    run_scraping(urls, custom_path)
    
    # Update status when finished
    status_label.config(text="Scraping and summarization completed!", fg="green")
    messagebox.showinfo("Info", "Scraping and summarization process completed!")

# Create the GUI window
root = tk.Tk()
root.title("Web Scraping Tool")

# Add a status label to the GUI window
status_label = tk.Label(root, text="", fg="blue")
status_label.pack(pady=10)

# Create and place the URL input label and entry
tk.Label(root, text="Enter the URL(s) to scrape data from (one per line):").pack(pady=10)
url_entry = tk.Text(root, height=5, width=50)
url_entry.pack(pady=5)

# Create and place the storage path label and entry
tk.Label(root, text="Enter the path for storing the file after scraping:").pack(pady=10)
path_entry = tk.Entry(root, width=50)
path_entry.pack(pady=5)

# Create format selection dropdown
format_var = tk.StringVar(value="Excel")  # Default format
format_label = tk.Label(root, text="Select file format:")
format_label.pack(pady=10)
format_dropdown = tk.OptionMenu(root, format_var, "Excel", "CSV", "Word")
format_dropdown.pack(pady=5)

# Create progress bar
progress_bar = ttk.Progressbar(root, length=200, mode='determinate')
progress_bar.pack(pady=20)

# Create and place the Start button
start_button = tk.Button(root, text="Start Scraping", command=start_scraping)
start_button.pack(pady=20)

# Start the GUI loop
root.mainloop()
